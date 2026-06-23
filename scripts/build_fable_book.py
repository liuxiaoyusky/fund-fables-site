"""
Build combined fables book (Phase 8 of pdf-to-study-program).

Reads a chapter-organised fable corpus under `<source>/`, scans every chapter
folder in sorted order and every fable `.md` inside each chapter in numbered
order, and emits three artefacts under `<out-dir>/`:

    <book-id>-book.md     merged single-file Markdown
    <book-id>-book.docx   Word document with CJK font, TOC, page numbers
    <book-id>-book.pdf    PDF with CJK font, TOC, page numbers

Usage:
    # Single book, all args on CLI
    python3 build_fable_book.py \\
        --source docs/fables/fund-fables \\
        --book-id fund-fables \\
        --title "证券投资基金 · 寓言四百六十四" \\
        --subtitle "基金从业 · 科目二 · 寓言化" \\
        --author "原作：中国证券投资基金业协会" \\
        --compiler "寓言整理：王恩培 · 等"

    # All books from a config file
    python3 build_fable_book.py --config scripts/books.toml --all

The script expects each fable `.md` to follow the standard template:

    # <heading>
    ## <emoji> 寓言故事 —— 《<4–6字故事名>》
    <story body>
    ---
    **📖 原文定义**
    > quote lines...
    **💡 对应点**
    | 故事元素 | 概念对应 |
    |---|---|
    | row 1   | row 1   |
    > 📝 来源：...
"""

from __future__ import annotations

import argparse
import datetime
import os
import re
import sys
import tomllib
from pathlib import Path

# ---------------------------------------------------------------------------
# Defaults
# ---------------------------------------------------------------------------

DEFAULT_OUT_ROOT = Path("docs/downloads")
TODAY = datetime.date.today().isoformat()

# CJK font candidates — first one that registers wins for PDF.
# On macOS, STHeiti is preferred. On Linux (CI), the runner installs
# `fonts-noto-cjk`, which lives under /usr/share/fonts/.
CJK_FONT_CANDIDATES = [
    # macOS (local builds)
    ("STHeiti", "/System/Library/Fonts/STHeiti Medium.ttc", 0),
    ("STHeitiLight", "/System/Library/Fonts/STHeiti Light.ttc", 0),
    ("PingFang", "/Library/Fonts/Arial Unicode.ttf", 0),
    # Linux (GitHub Actions ubuntu-latest + fonts-noto-cjk)
    ("NotoSansCJK", "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", 0),
    ("NotoSansCJKsc", "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf", 0),
    ("WenQuanYi", "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", 0),
]


# ---------------------------------------------------------------------------
# Chapter / fable discovery
# ---------------------------------------------------------------------------

def natural_key(s: str):
    """Sort key that splits out leading numbers so '02-...' < '10-...'."""
    parts = re.split(r"(\d+)", s)
    return [int(p) if p.isdigit() else p.lower() for p in parts]


def chapter_display_title(slug: str) -> str:
    """Turn `01-证券投资基金概述` -> `证券投资基金概述`."""
    # Strip leading NN- prefix.
    cleaned = re.sub(r"^\d+-", "", slug)
    return cleaned.strip() or slug


def discover_chapters(source_dir: Path) -> list[tuple[Path, str]]:
    """Return [(chapter_dir, display_title), ...] in sorted order.

    A chapter is any immediate subdirectory of `source_dir` whose name looks
    like `NN-something` (starts with digits). The display title is the slug
    with the leading `NN-` stripped.
    """
    if not source_dir.is_dir():
        raise FileNotFoundError(f"source dir not found: {source_dir}")
    chapters = []
    for p in sorted(source_dir.iterdir(), key=lambda x: natural_key(x.name)):
        if not p.is_dir():
            continue
        if not re.match(r"^\d+[-_]", p.name):
            continue
        chapters.append((p, chapter_display_title(p.name)))
    if not chapters:
        raise RuntimeError(f"no NN-* chapter dirs found under {source_dir}")
    return chapters


def discover_fables(chapter_dir: Path) -> list[Path]:
    """Recursively find all `.md` fable files under a chapter, in numeric order.

    Files may live directly in the chapter (`chapter_dir/01-foo.md`) or one
    level deeper (`chapter_dir/01-foo/01-bar.md`). Sort by the leading
    `<NN>-` segment of the filename (natural numeric).
    """
    md_files = list(chapter_dir.rglob("*.md"))
    if not md_files:
        return []

    def sort_key(p: Path) -> tuple[int, str]:
        m = re.match(r"^(\d+)-", p.name)
        n = int(m.group(1)) if m else 0
        # Tie-break by relative path so subfolder files come after their
        # siblings in deterministic order.
        return (n, str(p.relative_to(chapter_dir)))

    return sorted(md_files, key=sort_key)


# ---------------------------------------------------------------------------
# Fable parser (verbatim from 100-game build_fables_book.py, no behaviour
# change). See SKILL.md Phase 8 / Step 8.1.
# ---------------------------------------------------------------------------

def _section_marker(ln: str) -> str | None:
    """Match both '## 原文定义' (100-game) and '**📖 原文定义**' (基金书)."""
    s = ln.strip()
    # H2 form: `## 原文定义` or `## 📖 原文定义`
    m = re.match(r"^##\s+(?:📖\s*)?原文定义\s*$", s)
    if m:
        return "quote"
    m = re.match(r"^##\s+(?:💡\s*)?对应点\s*$", s)
    if m:
        return "table"
    # Bold form: `**📖 原文定义**` or `**原文定义**`
    m = re.match(r"^\*\*(?:📖\s*)?原文定义\*\*\s*$", s)
    if m:
        return "quote"
    m = re.match(r"^\*\*(?:💡\s*)?对应点\*\*\s*$", s)
    if m:
        return "table"
    return None


def parse_fable(path: Path) -> dict:
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()

    # Strip YAML frontmatter.
    if lines and lines[0].strip() == "---":
        for i in range(1, len(lines)):
            if lines[i].strip() == "---":
                lines = lines[i + 1:]
                break

    # Title = first H1.
    title = ""
    for ln in lines:
        m = re.match(r"^#\s+(.+?)\s*$", ln)
        if m:
            title = m.group(1).strip()
            break
    if not title:
        title = f"寓言 {path.stem}"

    i_quote = i_table = None
    for i, ln in enumerate(lines):
        kind = _section_marker(ln)
        if kind == "quote" and i_quote is None:
            i_quote = i
        elif kind == "table" and i_table is None:
            i_table = i
    if i_quote is None or i_table is None:
        raise ValueError(f"{path}: missing 原文定义 or 对应点 section")

    body_start = 0
    for i, ln in enumerate(lines):
        if re.match(r"^#\s+", ln):
            body_start = i + 1
            break
    body_lines = lines[body_start:i_quote]
    while body_lines and not body_lines[0].strip():
        body_lines.pop(0)
    while body_lines and not body_lines[-1].strip():
        body_lines.pop()

    source_line = ""
    for ln in lines:
        if "📝 来源" in ln:
            cleaned = re.sub(r"^>\s*", "", ln).strip()
            cleaned = re.sub(r"^📝\s*", "", cleaned).strip()
            source_line = cleaned
            break
    body_lines = [ln for ln in body_lines if "📝 来源" not in ln]
    body = "\n".join(body_lines).rstrip()

    quote_lines = []
    for ln in lines[i_quote + 1:i_table]:
        stripped = ln.strip()
        if not stripped:
            continue
        cleaned = re.sub(r"^>\s?", "", ln).rstrip()
        quote_lines.append(cleaned)

    table_rows = []
    for ln in lines[i_table + 1:]:
        s = ln.strip()
        if not s.startswith("|"):
            continue
        if s.count("|") < 3:
            continue
        cells = [c.strip() for c in s.strip("|").split("|")]
        if all(re.fullmatch(r"[-:\s]+", c) for c in cells):
            continue
        table_rows.append(cells)
    if table_rows and len(table_rows[0]) == 2 and table_rows[0][0] == "故事元素":
        table_rows = table_rows[1:]

    return {
        "title": title,
        "body": body,
        "quote_lines": quote_lines,
        "table_rows": table_rows,
        "source_line": source_line,
    }


# ---------------------------------------------------------------------------
# 1. Merged Markdown
# ---------------------------------------------------------------------------

def write_merged_markdown(
    out_path: Path,
    fables_by_chapter: list[tuple[str, str, list[tuple[int, dict, Path]]]],
    book_title: str,
    book_subtitle: str,
    author_original: str,
    attribution: str,
) -> int:
    total_fables = sum(len(v) for _, _, v in fables_by_chapter)

    out = []
    out.append(f"# {book_title}\n")
    out.append(f"## {book_subtitle}\n")
    out.append(f"{author_original}\n")
    out.append(f"{attribution}\n")
    out.append(f"生成日期：{TODAY}\n")
    out.append(f"寓言数量：{total_fables} 篇 · 共 {len(fables_by_chapter)} 章\n")
    out.append("")

    out.append("## 目录\n")
    for idx, (_slug, en_title, items) in enumerate(fables_by_chapter, 1):
        n = len(items)
        out.append(f"- 第 {idx} 章 {en_title} （{n} 篇）")
    out.append("")

    out.append("---\n")

    for idx, (_slug, en_title, items) in enumerate(fables_by_chapter, 1):
        out.append(f"# 第 {idx} 章 {en_title}\n")
        n = len(items)
        out.append(f"> 本章包含 {n} 篇寓言\n")
        out.append("")

        for _sec_no, fable, _src in items:
            out.append(f"## {fable['title']}\n")
            out.append(fable["body"])
            out.append("")
            out.append("### 原文定义\n")
            for q in fable["quote_lines"]:
                out.append(f"> {q}")
            out.append("")
            out.append("### 对应点\n")
            table_rows = fable["table_rows"]
            if table_rows and len(table_rows[0]) == 2 and table_rows[0][0] == "故事元素":
                table_rows = table_rows[1:]
            if table_rows:
                out.append("| 故事元素 | 概念对应 |")
                out.append("|---|---|")
                for row in table_rows:
                    cells = (row + ["", ""])[:2]
                    out.append(f"| {cells[0]} | {cells[1]} |")
            else:
                out.append("_(无)_")
            out.append("")
            if fable["source_line"]:
                out.append(f"📝 {fable['source_line']}")
                out.append("")
            out.append("---\n")

    md_text = "\n".join(out)
    out_path.write_text(md_text, encoding="utf-8")
    return md_text.count("\n") + 1


# ---------------------------------------------------------------------------
# 2. DOCX
# ---------------------------------------------------------------------------

def _set_cjk(run, name: str, size: float, bold: bool = False,
             italic: bool = False, color=None):
    from docx.shared import Pt, RGBColor  # noqa
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def _add_page_number_footer(section, cjk_name: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.text = "PAGE   \\* MERGEFORMAT"
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
    _set_cjk(run, cjk_name, size=9, color=None)


def build_docx(
    out_path: Path,
    fables_by_chapter: list[tuple[str, str, list[tuple[int, dict, Path]]]],
    book_title: str,
    book_subtitle: str,
    author_original: str,
    attribution: str,
    cjk_name: str = "PingFang SC",
):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    total_fables = sum(len(v) for _, _, v in fables_by_chapter)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = cjk_name
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), cjk_name)
    rFonts.set(qn("w:ascii"), cjk_name)
    rFonts.set(qn("w:hAnsi"), cjk_name)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ----- Title page -----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run(book_title)
    _set_cjk(r, cjk_name, size=28, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(book_subtitle)
    _set_cjk(r, cjk_name, size=14, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run(author_original)
    _set_cjk(r, cjk_name, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(attribution)
    _set_cjk(r, cjk_name, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run(f"生成日期：{TODAY}")
    _set_cjk(r, cjk_name, size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"寓言数量：{total_fables} 篇 · 共 {len(fables_by_chapter)} 章")
    _set_cjk(r, cjk_name, size=10)

    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    # ----- TOC -----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("目录")
    _set_cjk(r, cjk_name, size=22, bold=True)
    p.paragraph_format.space_after = Pt(18)

    for idx, (_slug, en_title, items) in enumerate(fables_by_chapter, 1):
        n = len(items)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"第 {idx} 章　{en_title}　（{n} 篇）")
        _set_cjk(r, cjk_name, size=12)

    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    # ----- Chapters -----
    for idx, (_slug, en_title, items) in enumerate(fables_by_chapter, 1):
        n = len(items)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(180)
        r = p.add_run(f"第 {idx} 章")
        _set_cjk(r, cjk_name, size=18, bold=False)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        r = p.add_run(en_title)
        _set_cjk(r, cjk_name, size=24, bold=True)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(40)
        r = p.add_run(f"本章包含 {n} 篇寓言")
        _set_cjk(r, cjk_name, size=12, italic=True)

        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

        for _sec_no, fable, _src in items:
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(6)
            h.paragraph_format.space_after = Pt(12)
            h.paragraph_format.keep_with_next = True
            r = h.add_run(fable["title"])
            _set_cjk(r, cjk_name, size=18, bold=True)
            hPr = h._element.get_or_add_pPr()
            h_rPr = OxmlElement("w:rPr")
            h_rFonts = OxmlElement("w:rFonts")
            h_rFonts.set(qn("w:eastAsia"), cjk_name)
            h_rFonts.set(qn("w:ascii"), cjk_name)
            h_rFonts.set(qn("w:hAnsi"), cjk_name)
            h_rPr.append(h_rFonts)
            hPr.append(h_rPr)

            for para in re.split(r"\n\s*\n", fable["body"]):
                para = para.strip()
                if not para:
                    continue
                if re.fullmatch(r"-{3,}", para):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run("· · ·")
                    _set_cjk(r, cjk_name, size=10, color=RGBColor(0x99, 0x99, 0x99))
                    continue
                if para.startswith("## "):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(8)
                    p.paragraph_format.space_after = Pt(4)
                    r = p.add_run(para[3:].strip())
                    _set_cjk(r, cjk_name, size=12, bold=True, italic=True)
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0.74)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.5
                lines = para.split("\n")
                for j, ln in enumerate(lines):
                    if j > 0:
                        p.add_run().add_break()
                    r = p.add_run(ln)
                    _set_cjk(r, cjk_name, size=10.5)

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)

            h = doc.add_heading(level=3)
            r = h.add_run("原文定义")
            _set_cjk(r, cjk_name, size=13, bold=True)
            hPr = h._element.get_or_add_pPr()
            h_rPr = OxmlElement("w:rPr")
            h_rFonts = OxmlElement("w:rFonts")
            h_rFonts.set(qn("w:eastAsia"), cjk_name)
            h_rFonts.set(qn("w:ascii"), cjk_name)
            h_rFonts.set(qn("w:hAnsi"), cjk_name)
            h_rPr.append(h_rFonts)
            hPr.append(h_rPr)

            for q in fable["quote_lines"]:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.8)
                p.paragraph_format.right_indent = Cm(0.8)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.4
                r = p.add_run(q)
                _set_cjk(r, cjk_name, size=10.5, italic=True,
                         color=RGBColor(0x40, 0x40, 0x40))

            h = doc.add_heading(level=3)
            r = h.add_run("对应点")
            _set_cjk(r, cjk_name, size=13, bold=True)
            hPr = h._element.get_or_add_pPr()
            h_rPr = OxmlElement("w:rPr")
            h_rFonts = OxmlElement("w:rFonts")
            h_rFonts.set(qn("w:eastAsia"), cjk_name)
            h_rFonts.set(qn("w:ascii"), cjk_name)
            h_rFonts.set(qn("w:hAnsi"), cjk_name)
            h_rPr.append(h_rFonts)
            hPr.append(h_rPr)

            if fable["table_rows"]:
                tbl = doc.add_table(rows=1 + len(fable["table_rows"]), cols=2)
                tbl.style = "Table Grid"
                widths = [Cm(7.5), Cm(8.0)]
                hdr = tbl.rows[0].cells
                hdr[0].text = ""
                hdr[1].text = ""
                for cell, text in zip(hdr, ["故事元素", "概念对应"]):
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(text)
                    _set_cjk(r, cjk_name, size=10.5, bold=True)
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "D9D9D9")
                    tcPr.append(shd)
                for ri, row in enumerate(fable["table_rows"], start=1):
                    cells = (row + ["", ""])[:2]
                    for ci, val in enumerate(cells):
                        cell = tbl.rows[ri].cells[ci]
                        p = cell.paragraphs[0]
                        p.paragraph_format.line_spacing = 1.3
                        r = p.add_run(val)
                        _set_cjk(r, cjk_name, size=10)
                    tbl.rows[ri].height = Cm(0.6)
                for row in tbl.rows:
                    for ci, w in enumerate(widths):
                        row.cells[ci].width = w
            else:
                p = doc.add_paragraph()
                r = p.add_run("（无）")
                _set_cjk(r, cjk_name, size=10, italic=True,
                         color=RGBColor(0x80, 0x80, 0x80))

            if fable["source_line"]:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(8)
                r = p.add_run(f"📝 {fable['source_line']}")
                _set_cjk(r, cjk_name, size=9, color=RGBColor(0x80, 0x80, 0x80))

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(18)
            r = p.add_run("· · ·")
            _set_cjk(r, cjk_name, size=10, color=RGBColor(0x99, 0x99, 0x99))

    for sec in doc.sections:
        _add_page_number_footer(sec, cjk_name)

    doc.save(str(out_path))


# ---------------------------------------------------------------------------
# 3. PDF
# ---------------------------------------------------------------------------

def _register_cjk_font() -> str:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    for name, path, sub_idx in CJK_FONT_CANDIDATES:
        if not os.path.exists(path):
            continue
        try:
            pdfmetrics.registerFont(TTFont(name, path, subfontIndex=sub_idx))
            print(f"[pdf] using CJK font: {name} from {path}")
            return name
        except Exception as e:
            print(f"[pdf] failed to register {name} from {path}: {e}")
    raise RuntimeError("No usable CJK font found on this system.")


def _escape_xml(s: str) -> str:
    return (s.replace("&", "&amp;")
             .replace("<", "&lt;")
             .replace(">", "&gt;"))


def build_pdf(
    out_path: Path,
    fables_by_chapter: list[tuple[str, str, list[tuple[int, dict, Path]]]],
    book_title: str,
    book_subtitle: str,
    author_original: str,
    attribution: str,
    cjk_name: str,
):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib import colors
    from reportlab.platypus import (BaseDocTemplate, PageTemplate, Frame,
                                    Paragraph, Spacer, PageBreak, Table,
                                    TableStyle, NextPageTemplate)

    total_fables = sum(len(v) for _, _, v in fables_by_chapter)
    page_w, page_h = A4

    def cover_page(canv, doc):
        canv.saveState()
        canv.setFont(cjk_name, 9)
        canv.setFillColor(colors.grey)
        canv.restoreState()

    def content_page(canv, doc):
        canv.saveState()
        canv.setFont(cjk_name, 9)
        canv.setFillColor(colors.grey)
        canv.drawCentredString(page_w / 2, 1.2 * cm, str(canv.getPageNumber()))
        canv.restoreState()

    doc = BaseDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
        title=book_title,
        author=attribution,
    )

    frame_cover = Frame(0, 0, page_w, page_h, id="cover",
                        leftPadding=2.5 * cm, rightPadding=2.5 * cm,
                        topPadding=2.5 * cm, bottomPadding=2.5 * cm)
    frame_content = Frame(2.5 * cm, 2.5 * cm, page_w - 5 * cm, page_h - 5 * cm,
                          id="content", leftPadding=0, rightPadding=0,
                          topPadding=0, bottomPadding=0)

    doc.addPageTemplates([
        PageTemplate(id="cover", frames=[frame_cover], onPage=cover_page),
        PageTemplate(id="content", frames=[frame_content], onPage=content_page),
    ])

    s_title = ParagraphStyle("title", fontName=cjk_name, fontSize=32, leading=44,
                             alignment=TA_CENTER, spaceBefore=120, spaceAfter=12)
    s_subtitle = ParagraphStyle("subtitle", fontName=cjk_name, fontSize=15, leading=22,
                                alignment=TA_CENTER, spaceAfter=40,
                                textColor=colors.HexColor("#444444"))
    s_meta = ParagraphStyle("meta", fontName=cjk_name, fontSize=11, leading=18,
                            alignment=TA_CENTER, spaceAfter=4)
    s_meta_lg = ParagraphStyle("metalarge", fontName=cjk_name, fontSize=12, leading=20,
                               alignment=TA_CENTER, spaceAfter=6)
    s_toc_h = ParagraphStyle("toch", fontName=cjk_name, fontSize=22, leading=30,
                             spaceBefore=0, spaceAfter=20)
    s_toc_item = ParagraphStyle("tocitem", fontName=cjk_name, fontSize=12, leading=22,
                                spaceAfter=6)
    s_chap_num = ParagraphStyle("chapnum", fontName=cjk_name, fontSize=18, leading=28,
                                alignment=TA_CENTER, spaceBefore=180, spaceAfter=12)
    s_chap_title = ParagraphStyle("chaptitle", fontName=cjk_name, fontSize=26, leading=40,
                                  alignment=TA_CENTER, spaceAfter=30, bold=True)
    s_chap_meta = ParagraphStyle("chapmeta", fontName=cjk_name, fontSize=12, leading=20,
                                 alignment=TA_CENTER, spaceAfter=0,
                                 textColor=colors.HexColor("#666666"))
    s_fable_title = ParagraphStyle("fabletitle", fontName=cjk_name, fontSize=18, leading=28,
                                   spaceBefore=8, spaceAfter=12, bold=True)
    s_body = ParagraphStyle("body", fontName=cjk_name, fontSize=10.5, leading=18,
                            firstLineIndent=21, spaceAfter=4, alignment=TA_LEFT)
    s_h2_inline = ParagraphStyle("h2inline", fontName=cjk_name, fontSize=12, leading=20,
                                 spaceBefore=10, spaceAfter=6, bold=True, italic=True)
    s_sep = ParagraphStyle("sep", fontName=cjk_name, fontSize=10, leading=14,
                           alignment=TA_CENTER, spaceBefore=4, spaceAfter=4,
                           textColor=colors.HexColor("#999999"))
    s_h3 = ParagraphStyle("h3", fontName=cjk_name, fontSize=13, leading=22,
                          spaceBefore=12, spaceAfter=6, bold=True)
    s_quote = ParagraphStyle("quote", fontName=cjk_name, fontSize=10.5, leading=18,
                             leftIndent=20, rightIndent=20, spaceAfter=2,
                             textColor=colors.HexColor("#404040"), italic=True)
    s_source = ParagraphStyle("source", fontName=cjk_name, fontSize=9, leading=14,
                              alignment=TA_CENTER, spaceBefore=8, spaceAfter=4,
                              textColor=colors.HexColor("#808080"))

    story = []

    story.append(Paragraph(_escape_xml(book_title), s_title))
    story.append(Paragraph(_escape_xml(book_subtitle), s_subtitle))
    story.append(Spacer(1, 40))
    story.append(Paragraph(_escape_xml(author_original), s_meta_lg))
    story.append(Paragraph(_escape_xml(attribution), s_meta_lg))
    story.append(Spacer(1, 60))
    story.append(Paragraph(f"生成日期：{_escape_xml(TODAY)}", s_meta))
    story.append(Paragraph(f"寓言数量：{total_fables} 篇 · 共 {len(fables_by_chapter)} 章", s_meta))

    story.append(NextPageTemplate("content"))
    story.append(PageBreak())

    story.append(Paragraph("目录", s_toc_h))
    for idx, (_slug, en_title, items) in enumerate(fables_by_chapter, 1):
        n = len(items)
        story.append(Paragraph(
            f"第 {idx} 章　{_escape_xml(en_title)}　（{n} 篇）", s_toc_item))
    story.append(PageBreak())

    for idx, (_slug, en_title, items) in enumerate(fables_by_chapter, 1):
        n = len(items)
        story.append(Spacer(1, 60))
        story.append(Paragraph(f"第 {idx} 章", s_chap_num))
        story.append(Paragraph(_escape_xml(en_title), s_chap_title))
        story.append(Spacer(1, 30))
        story.append(Paragraph(f"本章包含 {n} 篇寓言", s_chap_meta))
        story.append(PageBreak())

        for _sec_no, fable, _src in items:
            story.append(Paragraph(_escape_xml(fable["title"]), s_fable_title))
            for para in re.split(r"\n\s*\n", fable["body"]):
                para = para.strip()
                if not para:
                    continue
                if re.fullmatch(r"-{3,}", para):
                    story.append(Paragraph("· · ·", s_sep))
                    continue
                if para.startswith("## "):
                    story.append(Paragraph(_escape_xml(para[3:].strip()), s_h2_inline))
                    continue
                story.append(Paragraph(_escape_xml(para).replace("\n", "<br/>"), s_body))

            story.append(Spacer(1, 4))
            story.append(Paragraph("原文定义", s_h3))
            quote_html = "<br/>".join(_escape_xml(q) for q in fable["quote_lines"])
            if quote_html:
                story.append(Paragraph(quote_html, s_quote))

            story.append(Paragraph("对应点", s_h3))
            if fable["table_rows"]:
                cell_style = ParagraphStyle(
                    "tablecell", fontName=cjk_name, fontSize=9.5, leading=14,
                    alignment=TA_LEFT, spaceBefore=0, spaceAfter=0,
                )
                cell_style_h = ParagraphStyle(
                    "tablecellh", parent=cell_style, fontSize=10, leading=15,
                    alignment=TA_CENTER,
                )
                data = [[Paragraph("故事元素", cell_style_h),
                         Paragraph("概念对应", cell_style_h)]]
                for row in fable["table_rows"]:
                    cells = (row + ["", ""])[:2]
                    data.append([
                        Paragraph(_escape_xml(cells[0]), cell_style),
                        Paragraph(_escape_xml(cells[1]), cell_style),
                    ])
                tbl = Table(data, colWidths=[7.5 * cm, 8.0 * cm])
                tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9D9D9")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#999999")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]))
                story.append(tbl)
            else:
                story.append(Paragraph("（无）", s_quote))

            if fable["source_line"]:
                story.append(Paragraph(f"📝 {_escape_xml(fable['source_line'])}", s_source))
            story.append(Paragraph("· · ·", s_sep))
            story.append(Spacer(1, 12))

    doc.build(story)


# ---------------------------------------------------------------------------
# Orchestrator
# ---------------------------------------------------------------------------

def load_books_config(config_path: Path) -> list[dict]:
    with open(config_path, "rb") as f:
        cfg = tomllib.load(f)
    books = []
    for entry in cfg.get("books", []):
        for required in ("id", "source", "title", "subtitle", "author", "compiler"):
            if required not in entry:
                raise ValueError(f"books.toml entry missing `{required}`: {entry}")
        books.append(entry)
    if not books:
        raise ValueError("no `[[books]]` entries found in config")
    return books


def build_one_book(spec: dict) -> dict:
    source_dir = Path(spec["source"])
    book_id = spec["id"]
    out_dir = Path(spec.get("out_dir", DEFAULT_OUT_ROOT)) / book_id
    out_dir.mkdir(parents=True, exist_ok=True)

    print(f"\n[book] id={book_id}  source={source_dir}  out={out_dir}")

    chapters = discover_chapters(source_dir)
    fables_by_chapter: list[tuple[str, str, list[tuple[int, dict, Path]]]] = []
    for chap_dir, chap_title in chapters:
        files = discover_fables(chap_dir)
        if not files:
            print(f"  WARN: chapter '{chap_dir.name}' has no .md files, skipping")
            continue
        items: list[tuple[int, dict, Path]] = []
        for p in files:
            try:
                fable = parse_fable(p)
                m = re.match(r"^(\d+)-", p.name)
                sec_no = int(m.group(1)) if m else len(items) + 1
                items.append((sec_no, fable, p))
            except Exception as e:
                print(f"  WARN: skipping {p.relative_to(source_dir)}: {e}")
        items.sort(key=lambda x: x[0])
        fables_by_chapter.append((chap_dir.name, chap_title, items))
        print(f"  chapter '{chap_dir.name}' ({chap_title}): {len(items)} fables")

    out_md = out_dir / f"{book_id}-book.md"
    out_docx = out_dir / f"{book_id}-book.docx"
    out_pdf = out_dir / f"{book_id}-book.pdf"

    print(f"[md] writing merged markdown -> {out_md}")
    md_lines = write_merged_markdown(
        out_md, fables_by_chapter,
        spec["title"], spec["subtitle"], spec["author"], spec["compiler"],
    )
    print(f"  -> {md_lines} lines, {out_md.stat().st_size} bytes")

    print(f"[docx] building Word document -> {out_docx}")
    build_docx(
        out_docx, fables_by_chapter,
        spec["title"], spec["subtitle"], spec["author"], spec["compiler"],
    )
    print(f"  -> {out_docx.stat().st_size} bytes")

    print(f"[pdf] building PDF -> {out_pdf}")
    cjk = _register_cjk_font()
    build_pdf(
        out_pdf, fables_by_chapter,
        spec["title"], spec["subtitle"], spec["author"], spec["compiler"],
        cjk_name=cjk,
    )
    print(f"  -> {out_pdf.stat().st_size} bytes")

    total = sum(len(v) for _, _, v in fables_by_chapter)
    return {"id": book_id, "chapters": len(fables_by_chapter), "fables": total,
            "md": out_md, "docx": out_docx, "pdf": out_pdf}


def main() -> int:
    p = argparse.ArgumentParser(description=__doc__.split("\n\n")[0])
    p.add_argument("--source", help="fables root dir (single-book mode)")
    p.add_argument("--book-id", help="book slug (single-book mode)")
    p.add_argument("--title", help="cover title")
    p.add_argument("--subtitle", help="cover subtitle")
    p.add_argument("--author", help="original author line")
    p.add_argument("--compiler", help="fable compiler line")
    p.add_argument("--out-dir", help="output root dir (default: docs/downloads)")
    p.add_argument("--config", help="path to books.toml (multi-book mode)")
    p.add_argument("--all", action="store_true",
                   help="with --config: build every book in the config")
    args = p.parse_args()

    if args.config:
        if not args.all:
            p.error("--config requires --all")
        specs = load_books_config(Path(args.config))
    else:
        if not (args.source and args.book_id and args.title and args.subtitle
                and args.author and args.compiler):
            p.error("single-book mode requires --source, --book-id, --title, "
                    "--subtitle, --author, --compiler")
        specs = [{
            "id": args.book_id,
            "source": args.source,
            "title": args.title,
            "subtitle": args.subtitle,
            "author": args.author,
            "compiler": args.compiler,
            "out_dir": args.out_dir or str(DEFAULT_OUT_ROOT),
        }]

    summaries = []
    for spec in specs:
        summaries.append(build_one_book(spec))

    print("\n[book] summary:")
    for s in summaries:
        print(f"  {s['id']:<25} {s['chapters']:>3} ch · {s['fables']:>4} fables "
              f"· md={s['md'].stat().st_size//1024} KB "
              f"· docx={s['docx'].stat().st_size//1024} KB "
              f"· pdf={s['pdf'].stat().st_size//1024} KB")

    return 0


if __name__ == "__main__":
    sys.exit(main())